import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import tkinter.font as font
import os
import email
import json
import threading
import time
import re
import queue  # 用于线程间通信，解决GUI线程安全问题

class EnronEmailProcessor:
    def __init__(self, root):
        self.root = root
        self.root.title("Enron邮件处理工具")
        self.root.geometry("800x600")
        
        # 全局变量
        self.dataset_path = ""
        self.output_path = ""
        self.processed_count = 0
        self.total_count = 0
        self.queue = queue.Queue()  # 线程通信队列，存放GUI更新指令
        self.root.after(100, self.process_queue)  # 主线程轮询队列
        
        # 创建主框架
        self.main_frame = ttk.Frame(root, padding="20")
        self.main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题（兼容无SimHei字体的情况）
        self.title_label = ttk.Label(self.main_frame, text="Enron邮件处理工具", font=("SimHei", 16, "bold") if "SimHei" in font.families() else ("Arial", 16, "bold"))
        self.title_label.pack(pady=10)
        
        # 数据集选择部分
        self.dataset_frame = ttk.LabelFrame(self.main_frame, text="选择Enron数据集", padding="10")
        self.dataset_frame.pack(fill=tk.X, pady=10)
        
        self.dataset_path_var = tk.StringVar()
        ttk.Entry(self.dataset_frame, textvariable=self.dataset_path_var, width=60).pack(side=tk.LEFT, padx=5)
        ttk.Button(self.dataset_frame, text="浏览", command=self.select_dataset).pack(side=tk.LEFT, padx=5)
        
        # 输出文件选择部分
        self.output_frame = ttk.LabelFrame(self.main_frame, text="选择输出文件", padding="10")
        self.output_frame.pack(fill=tk.X, pady=10)
        
        self.output_path_var = tk.StringVar()
        ttk.Entry(self.output_frame, textvariable=self.output_path_var, width=60).pack(side=tk.LEFT, padx=5)
        ttk.Button(self.output_frame, text="浏览", command=self.select_output).pack(side=tk.LEFT, padx=5)
        
        # 处理按钮
        self.process_button = ttk.Button(self.main_frame, text="开始处理", command=self.start_processing)
        self.process_button.pack(pady=10)
        
        # 进度条
        self.progress_frame = ttk.LabelFrame(self.main_frame, text="处理进度", padding="10")
        self.progress_frame.pack(fill=tk.X, pady=10)
        
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(self.progress_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X, padx=5, pady=5)
        
        self.progress_label = ttk.Label(self.progress_frame, text="准备就绪")
        self.progress_label.pack(pady=5)
        
        # 状态显示
        self.status_frame = ttk.LabelFrame(self.main_frame, text="状态", padding="10")
        self.status_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        self.status_text = tk.Text(self.status_frame, height=10, wrap=tk.WORD)
        self.status_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.status_text.config(state=tk.DISABLED)
    
    def process_queue(self):
        """主线程处理队列中的GUI更新指令（线程安全）"""
        try:
            while True:
                task = self.queue.get_nowait()
                if task["type"] == "log":
                    self._safe_log(task["message"])
                elif task["type"] == "progress":
                    self.progress_var.set(task["value"])
                    self.progress_label.config(text=task["text"])
                elif task["type"] == "finish":
                    messagebox.showinfo("完成", task["message"])
                elif task["type"] == "error":
                    messagebox.showerror("错误", task["message"])
        except queue.Empty:
            pass
        finally:
            self.root.after(100, self.process_queue)  # 循环轮询
    
    def _safe_log(self, message):
        """线程安全的日志写入"""
        self.status_text.config(state=tk.NORMAL)
        self.status_text.insert(tk.END, f"{time.strftime('%Y-%m-%d %H:%M:%S')} - {message}\n")
        self.status_text.see(tk.END)
        self.status_text.config(state=tk.DISABLED)
    
    def log(self, message):
        """向队列添加日志指令"""
        self.queue.put({"type": "log", "message": message})
    
    def update_progress(self, value, text):
        """向队列添加进度更新指令"""
        self.queue.put({"type": "progress", "value": value, "text": text})
    
    def show_finish(self, message):
        """向队列添加完成提示指令"""
        self.queue.put({"type": "finish", "message": message})
    
    def show_error(self, message):
        """向队列添加错误提示指令"""
        self.queue.put({"type": "error", "message": message})
    
    def select_dataset(self):
        """选择Enron数据集目录"""
        path = filedialog.askdirectory(title="选择Enron数据集根目录")
        if path:
            self.dataset_path_var.set(path)
            self.dataset_path = path
            self.log(f"选择数据集目录: {path}")
    
    def select_output(self):
        """选择输出Word文档"""
        path = filedialog.asksaveasfilename(
            title="保存Word文档",
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self.output_path_var.set(path)
            self.output_path = path
            self.log(f"选择输出文件: {path}")
    
    def extract_email_content(self, email_path):
        """提取邮件的纯文本内容（修复编码问题）"""
        # 跳过空文件/隐藏文件
        if os.path.getsize(email_path) == 0 or os.path.basename(email_path).startswith('.'):
            return None
        
        try:
            # Enron数据集优先用latin-1编码读取（兼容utf-8）
            with open(email_path, 'r', encoding='latin-1', errors='ignore') as f:
                content = f.read()
            
            # 尝试解析为邮件
            try:
                msg = email.message_from_string(content)
                
                # 提取发件人、收件人、主题（清理空白字符）
                from_ = re.sub(r'\s+', ' ', msg.get('From', '')).strip()
                to = re.sub(r'\s+', ' ', msg.get('To', '')).strip()
                subject = re.sub(r'\s+', ' ', msg.get('Subject', '')).strip()
                
                # 提取纯文本内容
                email_content = ""
                if msg.is_multipart():
                    for part in msg.walk():
                        if part.get_content_type() == 'text/plain':
                            payload = part.get_payload(decode=True)
                            if payload:
                                # 兼容多种编码解码
                                for encoding in ['utf-8', 'latin-1', 'ascii']:
                                    try:
                                        email_content += payload.decode(encoding, errors='ignore')
                                        break
                                    except:
                                        continue
                else:
                    if msg.get_content_type() == 'text/plain':
                        payload = msg.get_payload(decode=True)
                        if payload:
                            for encoding in ['utf-8', 'latin-1', 'ascii']:
                                try:
                                    email_content = payload.decode(encoding, errors='ignore')
                                    break
                                except:
                                    continue
                
                # 清理内容
                email_content = self.clean_content(email_content)
                
                return {
                    'file_path': email_path,
                    'from': from_,
                    'to': to,
                    'subject': subject,
                    'content': email_content
                }
            except Exception as e:
                self.log(f"解析邮件失败 {email_path}，使用原始内容: {str(e)[:50]}")
                return {
                    'file_path': email_path,
                    'from': '',
                    'to': '',
                    'subject': '',
                    'content': self.clean_content(content)
                }
        except PermissionError:
            self.log(f"权限不足，无法读取: {email_path}")
            return None
        except Exception as e:
            self.log(f"处理邮件失败 {email_path}: {str(e)[:50]}")
            return None
    
    def clean_content(self, content):
        """优化内容清理逻辑"""
        # 移除HTML标签
        content = re.sub(r'<[^>]+>', '', content)
        # 移除邮件引用标记（多行）
        content = re.sub(r'^>.*?$', '', content, flags=re.MULTILINE)
        # 移除多余空白（包括换行/制表符）
        content = re.sub(r'\s+', ' ', content).strip()
        # 移除不可见字符
        content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', content)
        return content
    
    def find_email_files(self, directory):
        """优化邮件文件筛选"""
        email_files = []
        # 避免访问系统目录
        restricted_paths = ["AppData", "Local Settings", "Program Files", "Windows", "System32"]
        
        # 添加调试信息
        self.log(f"开始扫描目录: {directory}")
        
        try:
            total_scanned = 0
            for root, dirs, files in os.walk(directory):
                # 跳过隐藏目录和系统目录
                dirs[:] = [d for d in dirs if not d.startswith('.') and d not in restricted_paths]
                
                # 检查当前目录是否在受限路径中
                current_path = os.path.normpath(root)
                if any(restricted in current_path for restricted in restricted_paths):
                    continue
                
                # 记录当前扫描的目录
                self.log(f"扫描目录: {root}, 发现 {len(files)} 个文件")
                
                for file in files:
                    total_scanned += 1
                    # 跳过隐藏文件
                    if file.startswith('.'):
                        continue
                    
                    # 极其宽松的文件筛选：只排除明显不是邮件的文件
                    # Enron数据集的邮件文件通常没有扩展名，或者可能有.txt扩展名
                    exclude_extensions = ('.zip', '.rar', '.7z', '.pdf', '.doc', '.docx', '.jpg', '.jpeg', '.png', '.gif', '.db', '.dat', '.exe', '.py', '.csv', '.xlsx')
                    if file.endswith(exclude_extensions):
                        continue
                    
                    # 特别处理Enron数据集的典型文件名格式，如"1.", "2."等
                    if '.' in file and not any(file.endswith(ext) for ext in ['.eml', '.txt']):
                        # 这可能是Enron邮件文件，如"1.", "2."等
                        pass
                    
                    file_path = os.path.join(root, file)
                    
                    try:
                        # 只处理存在且非空的文件
                        if os.path.exists(file_path) and os.path.isfile(file_path):
                            size = os.path.getsize(file_path)
                            if size > 0:
                                email_files.append(file_path)
                                # 每找到10个文件记录一次
                                if len(email_files) % 10 == 0:
                                    self.log(f"已找到 {len(email_files)} 个邮件文件")
                    except Exception as e:
                        # 记录错误但继续处理其他文件
                        self.log(f"处理文件 {file_path} 时出错: {str(e)}")
                        pass
            
            self.log(f"扫描完成，共扫描 {total_scanned} 个文件，找到 {len(email_files)} 个邮件文件")
        except Exception as e:
            # 捕获任何异常，确保程序不会崩溃
            self.log(f"扫描过程中出错: {str(e)}")
            pass
        
        return email_files
    
    def process_emails(self):
        """处理所有邮件（子线程执行）"""
        import os
        # 检查输入输出路径
        if not self.dataset_path or not os.path.exists(self.dataset_path):
            self.show_error("请选择有效的Enron数据集目录")
            return
        if not self.output_path:
            self.show_error("请选择输出文件")
            return
        
        # 查找所有邮件文件
        self.log("开始查找邮件文件...")
        email_files = self.find_email_files(self.dataset_path)
        self.total_count = len(email_files)
        self.processed_count = 0
        
        if self.total_count == 0:
            self.show_error("未找到有效邮件文件")
            return
        
        self.log(f"找到 {self.total_count} 个邮件文件")
        
        # 开始处理
        success_count = 0
        emails_data = []
        try:
            for i, email_file in enumerate(email_files):
                # 处理邮件
                email_data = self.extract_email_content(email_file)
                if email_data and email_data['content']:
                    emails_data.append(email_data)
                    success_count += 1
                
                # 更新进度（线程安全）
                self.processed_count = i + 1
                progress = (self.processed_count / self.total_count) * 100
                self.update_progress(progress, f"处理中... {self.processed_count}/{self.total_count}")
                
                # 每处理100个邮件更新一次日志
                if self.processed_count % 100 == 0:
                    self.log(f"已处理 {self.processed_count}/{self.total_count} 个邮件，成功提取 {success_count} 个有效内容")
            
            # 写入Word文档
            if emails_data:
                import os
                
                # 确保输出目录存在
                output_dir = os.path.dirname(self.output_path)
                if output_dir and not os.path.exists(output_dir):
                    try:
                        os.makedirs(output_dir)
                        self.log(f"创建输出目录: {output_dir}")
                    except Exception as e:
                        self.log(f"创建输出目录失败: {str(e)}")
                        self.show_error(f"创建输出目录失败: {str(e)}")
                        return
                
                # 确保python-docx库已安装
                try:
                    from docx import Document
                except ImportError:
                    self.log("python-docx库未安装，尝试自动安装...")
                    try:
                        import pip
                        pip.main(['install', 'python-docx'])
                        from docx import Document
                        self.log("python-docx库安装成功")
                    except Exception as e:
                        self.log(f"安装python-docx库失败: {str(e)}")
                        self.show_error(f"安装python-docx库失败: {str(e)}")
                        return
                
                # 创建Word文档
                doc = Document()
                doc.add_heading('Enron邮件数据集', 0)
                
                # 为每个邮件添加内容
                for i, email_data in enumerate(emails_data):
                    doc.add_heading(f'邮件 {i+1}', level=1)
                    doc.add_paragraph(f'文件路径: {email_data.get("file_path", "")}')
                    doc.add_paragraph(f'发件人: {email_data.get("from", "")}')
                    doc.add_paragraph(f'收件人: {email_data.get("to", "")}')
                    doc.add_paragraph(f'主题: {email_data.get("subject", "")}')
                    doc.add_heading('内容', level=2)
                    doc.add_paragraph(email_data.get("content", ""))
                    doc.add_page_break()
                
                # 尝试写入Word文档
                try:
                    doc.save(self.output_path)
                    self.log(f"成功写入 {len(emails_data)} 个邮件到Word文档")
                except PermissionError:
                    # 尝试使用临时文件
                    import tempfile
                    temp_path = tempfile.mktemp('.docx')
                    doc.save(temp_path)
                    
                    # 尝试复制到目标位置
                    try:
                        import shutil
                        shutil.copy2(temp_path, self.output_path)
                        os.unlink(temp_path)
                        self.log(f"成功写入 {len(emails_data)} 个邮件到Word文档（使用临时文件）")
                    except Exception as e:
                        self.log(f"写入Word文档失败: {str(e)}")
                        self.show_error(f"写入Word文档失败: {str(e)}")
                        return
            else:
                self.log("没有找到有效的邮件文件")
            
            self.update_progress(100, "处理完成")
            self.log(f"处理完成！共扫描 {self.processed_count} 个邮件，成功提取 {success_count} 个有效内容，输出到 {self.output_path}")
            self.show_finish(f"处理完成！共扫描 {self.processed_count} 个邮件，成功提取 {success_count} 个有效内容")
        except PermissionError:
            self.log("权限不足，无法写入输出文件")
            self.show_error("权限不足，无法写入输出文件（请检查文件是否被占用/路径是否有权限）")
        except Exception as e:
            self.log(f"处理过程中出错: {str(e)}")
            self.show_error(f"处理过程中出错: {str(e)}")
    
    def start_processing(self):
        """开始处理邮件（主线程触发）"""
        # 禁用按钮防止重复点击
        self.process_button.config(state=tk.DISABLED)
        # 重置进度
        self.update_progress(0, "准备处理...")
        self.log("开始处理Enron邮件...")
        
        # 启动子线程
        thread = threading.Thread(target=self.process_emails)
        thread.daemon = True
        thread.start()
        
        # 监听线程结束，恢复按钮
        def check_thread():
            if thread.is_alive():
                self.root.after(100, check_thread)
            else:
                self.process_button.config(state=tk.NORMAL)
        self.root.after(100, check_thread)

if __name__ == "__main__":
    root = tk.Tk()
    app = EnronEmailProcessor(root)
    root.mainloop()